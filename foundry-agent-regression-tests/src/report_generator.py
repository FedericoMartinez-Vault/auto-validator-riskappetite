"""Generate Markdown, Word, and chart outputs for segmented underwriting regression tests."""

from __future__ import annotations

import json
import re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

REPORT_DOCX_NAME = "AF-UW-RiskAppetite_Segmented_Underwriting_Test_Report.docx"
CHART_SPECS = [
    ("01_classification_distribution.png", "classification_distribution"),
    ("02_evaluation_distribution.png", "evaluation_distribution"),
    ("03_confidence_by_test.png", "confidence_by_test"),
    ("04_classification_by_guideline_track.png", "classification_by_guideline_track"),
    ("05_classification_by_outcome_segment.png", "classification_by_outcome_segment"),
    ("06_classification_by_territory_segment.png", "classification_by_territory_segment"),
    ("07_average_confidence_by_segment.png", "average_confidence_by_segment"),
    ("08_missing_information_frequency.png", "missing_information_frequency"),
]

OBSERVATION_KEYWORDS = (
    "moderate",
    "observation",
    "exception",
    "occupancy",
    "condo",
    "emerging",
    "needs ",
    "context",
    "dependency",
)


def _safe_str(value: Any) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    return str(value).strip()


def _parse_list_field(value: Any) -> list[str]:
    text = _safe_str(value)
    if not text:
        return []
    try:
        parsed = json.loads(text)
        if isinstance(parsed, list):
            return [str(item) for item in parsed if str(item).strip()]
    except json.JSONDecodeError:
        pass
    return [part.strip() for part in re.split(r"[;|]", text) if part.strip()]


def infer_evaluation_result(row: pd.Series) -> str:
    """Infer Pass / Pass with Observation / Fail / Needs Manual Review."""
    if _safe_str(row.get("error_message")):
        return "Fail"
    if _safe_str(row.get("parse_status")) == "failed":
        return "Fail"

    expected = _safe_str(row.get("expected_agent_focus")).lower()
    classification = _safe_str(row.get("agent_classification"))

    if "bad fit" in expected and classification == "Bad Fit":
        return "Pass"
    if any(keyword in expected for keyword in OBSERVATION_KEYWORDS):
        if classification in {"Moderate Fit", "Bad Fit", "Good Fit"}:
            return "Pass with Observation"
    if "good fit" in expected and classification == "Good Fit":
        return "Pass"
    if "bad fit" in expected and classification in {"Moderate Fit", "Good Fit"}:
        return "Fail"
    if "good fit" in expected and classification == "Bad Fit":
        return "Fail"
    if classification:
        return "Needs Manual Review"
    return "Needs Manual Review"


def short_finding(row: pd.Series) -> str:
    summary = _safe_str(row.get("summary"))
    if summary:
        return summary[:180] + ("..." if len(summary) > 180 else "")
    flags = _parse_list_field(row.get("risk_flags"))
    if flags:
        return flags[0][:180]
    return _safe_str(row.get("error_message")) or "No summary returned."


def enrich_results(df: pd.DataFrame) -> pd.DataFrame:
    enriched = df.copy()
    if "evaluation_result" not in enriched.columns:
        enriched["evaluation_result"] = enriched.apply(infer_evaluation_result, axis=1)
    if "short_finding" not in enriched.columns:
        enriched["short_finding"] = enriched.apply(short_finding, axis=1)
    enriched["confidence_score"] = pd.to_numeric(enriched["confidence_score"], errors="coerce")
    return enriched


def _count_table(df: pd.DataFrame, column: str) -> pd.DataFrame:
    if column not in df.columns:
        return pd.DataFrame(columns=[column, "count"])
    counts = (
        df[column]
        .fillna("Unknown")
        .replace("", "Unknown")
        .value_counts()
        .reset_index()
    )
    counts.columns = [column, "count"]
    return counts


def _save_current_figure(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout()
    plt.savefig(path, dpi=160, bbox_inches="tight")
    plt.close()


def generate_charts(df: pd.DataFrame, charts_dir: Path) -> dict[str, Path]:
    charts_dir.mkdir(parents=True, exist_ok=True)
    chart_paths: dict[str, Path] = {}

    classification_counts = df["agent_classification"].fillna("Unknown").replace("", "Unknown").value_counts()
    plt.figure(figsize=(8, 5))
    classification_counts.plot(kind="bar", color="#2E5AAC")
    plt.title("Classification Distribution")
    plt.xlabel("Classification")
    plt.ylabel("Count")
    plt.xticks(rotation=20, ha="right")
    path = charts_dir / "01_classification_distribution.png"
    _save_current_figure(path)
    chart_paths["classification_distribution"] = path

    evaluation_counts = df["evaluation_result"].fillna("Unknown").value_counts()
    plt.figure(figsize=(8, 5))
    evaluation_counts.plot(kind="bar", color="#3A7D44")
    plt.title("Evaluation Distribution")
    plt.xlabel("Evaluation Result")
    plt.ylabel("Count")
    plt.xticks(rotation=20, ha="right")
    path = charts_dir / "02_evaluation_distribution.png"
    _save_current_figure(path)
    chart_paths["evaluation_distribution"] = path

    plt.figure(figsize=(12, 5))
    plt.plot(range(len(df)), df["confidence_score"], marker="o", linewidth=1.5, color="#C47A2C")
    plt.title("Confidence Score by Test")
    plt.xlabel("Test Index")
    plt.ylabel("Confidence Score")
    plt.ylim(0, 1.05)
    plt.grid(axis="y", alpha=0.3)
    path = charts_dir / "03_confidence_by_test.png"
    _save_current_figure(path)
    chart_paths["confidence_by_test"] = path

    def grouped_classification(data: pd.DataFrame, segment_col: str, filename: str, key: str) -> None:
        if segment_col not in data.columns:
            return
        pivot = pd.crosstab(data[segment_col].fillna("Unknown"), data["agent_classification"].fillna("Unknown"))
        pivot.plot(kind="bar", stacked=True, figsize=(10, 5), colormap="tab20")
        plt.title(f"Classification by {segment_col.replace('_', ' ').title()}")
        plt.xlabel(segment_col.replace("_", " ").title())
        plt.ylabel("Count")
        plt.xticks(rotation=30, ha="right")
        plt.legend(title="Classification", bbox_to_anchor=(1.02, 1), loc="upper left")
        path = charts_dir / filename
        _save_current_figure(path)
        chart_paths[key] = path

    grouped_classification(
        df,
        "guideline_track",
        "04_classification_by_guideline_track.png",
        "classification_by_guideline_track",
    )
    grouped_classification(
        df,
        "outcome_segment",
        "05_classification_by_outcome_segment.png",
        "classification_by_outcome_segment",
    )
    grouped_classification(
        df,
        "territory_segment",
        "06_classification_by_territory_segment.png",
        "classification_by_territory_segment",
    )

    segment_conf = (
        df.groupby("guideline_track", dropna=False)["confidence_score"]
        .mean()
        .sort_values(ascending=False)
    )
    plt.figure(figsize=(8, 5))
    segment_conf.plot(kind="bar", color="#6C4F8D")
    plt.title("Average Confidence by Guideline Track")
    plt.xlabel("Guideline Track")
    plt.ylabel("Average Confidence")
    plt.xticks(rotation=15, ha="right")
    plt.ylim(0, 1.05)
    path = charts_dir / "07_average_confidence_by_segment.png"
    _save_current_figure(path)
    chart_paths["average_confidence_by_segment"] = path

    missing_counter: Counter[str] = Counter()
    for value in df.get("missing_information", pd.Series(dtype=str)):
        for item in _parse_list_field(value):
            missing_counter[item] += 1
    if missing_counter:
        items, counts = zip(*missing_counter.most_common(15))
    else:
        items, counts = (["No missing information flagged"], [len(df)])
    plt.figure(figsize=(10, 5))
    plt.barh(items[::-1], counts[::-1], color="#8B3A3A")
    plt.title("Missing Information Frequency")
    plt.xlabel("Frequency")
    path = charts_dir / "08_missing_information_frequency.png"
    _save_current_figure(path)
    chart_paths["missing_information_frequency"] = path

    return chart_paths


def _markdown_table(counts_df: pd.DataFrame) -> str:
    if counts_df.empty:
        return "_No data available._\n"
    header = f"| {counts_df.columns[0]} | count |\n| --- | ---: |\n"
    rows = "\n".join(f"| {row[0]} | {row[1]} |" for row in counts_df.itertuples(index=False))
    return header + rows + "\n"


def build_segment_analysis_markdown(df: pd.DataFrame) -> str:
    total = len(df)
    bad_fit = int((df["agent_classification"] == "Bad Fit").sum())
    moderate_fit = int((df["agent_classification"] == "Moderate Fit").sum())
    good_fit = int((df["agent_classification"] == "Good Fit").sum())
    avg_conf = df["confidence_score"].mean()
    parsed_ok = int(df["parse_status"].isin(["success", "recovered_json"]).sum())

    ves = df[df["guideline_track"].astype(str).str.contains("VES", na=False)]
    vre = df[df["guideline_track"].astype(str).str.contains("VRE", na=False)]
    issued_offered_bad = df[
        df["outcome_segment"].isin(["Issued", "Offered"]) & (df["agent_classification"] == "Bad Fit")
    ]

    lines = [
        "# Segment Analysis",
        "",
        f"_Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}_",
        "",
        "## Executive Snapshot",
        "",
        f"- Total tests run: **{total}**",
        f"- Parsed successfully: **{parsed_ok}**",
        f"- Bad Fit / Moderate Fit / Good Fit: **{bad_fit} / {moderate_fit} / {good_fit}**",
        f"- Average confidence: **{avg_conf:.4f}**" if pd.notna(avg_conf) else "- Average confidence: **N/A**",
        "",
        "This package supports segmented underwriting regression testing for the "
        "`AF-UW-RiskAppetite` agent. It is not a cybersecurity assessment.",
        "",
        "## VES / Non-Admitted",
        "",
        f"- Records tested: {len(ves)}",
        f"- Bad Fit classifications: {int((ves['agent_classification'] == 'Bad Fit').sum())}",
        "- The agent should select the VES guideline, apply closed coastal county and "
        "Harris <= 10 miles rules, and treat Coverage A / deductible issues as underwriting drivers.",
        "",
        "## VRE / Admitted",
        "",
        f"- Records tested: {len(vre)}",
        f"- Bad Fit classifications: {int((vre['agent_classification'] == 'Bad Fit').sum())}",
        "- The agent should select the VRE guideline, apply ineligible coastal tier / Harris "
        "<= 10 miles restrictions, and evaluate Coverage A below $2M correctly.",
        "",
        "## Outcome Segments",
        "",
        f"- Issued/Offered records classified as Bad Fit: **{len(issued_offered_bad)}**",
        "- This indicates the agent did not treat historical quote status as underwriting truth.",
        "",
    ]

    if not issued_offered_bad.empty:
        lines.append("### Examples")
        lines.append("")
        for _, row in issued_offered_bad.head(5).iterrows():
            lines.append(
                f"- `{row.get('quote_no')}` ({row.get('outcome_segment')}): "
                f"{_safe_str(row.get('short_finding'))}"
            )
        lines.append("")

    lines.extend(
        [
            "## Data Quality Handling",
            "",
            "- Missing company/program/location, zero Coverage A/TIV, and missing prior claims "
            "should be surfaced in `missing_information` and/or `risk_flags`.",
            "- Missing occupancy, roof, construction, plumbing, and electrical details should "
            "reduce confidence or trigger observation-level outcomes.",
            "",
            "## Instruction Compliance / Task Drift Resistance",
            "",
            "- Prompts required strict JSON underwriting output without markdown or citations.",
            "- `expected_agent_focus` was excluded from prompts to avoid biasing classifications.",
            "",
            "## Segment Count Tables",
            "",
        ]
    )

    for column in [
        "guideline_track",
        "outcome_segment",
        "territory_segment",
        "coverage_segment",
        "tiv_segment",
        "alarm_segment",
        "sprinkler_segment",
        "fire_protection_segment",
        "agent_classification",
        "evaluation_result",
    ]:
        lines.append(f"### {column}")
        lines.append("")
        lines.append(_markdown_table(_count_table(df, column)))
        lines.append("")

    return "\n".join(lines)


def _add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AF-UW-RiskAppetite Agent — Segmented Underwriting Test Results")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Texas VES / VRE HNW P&C Guideline-Based Evaluation"
    ).font.size = Pt(14)

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared.add_run("Prepared for: Vault Insurance").font.size = Pt(12)

    doc.add_paragraph("")
    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp.add_run(datetime.now().strftime("%B %d, %Y")).italic = True
    doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_counts_table(doc: Document, counts_df: pd.DataFrame, title: str) -> None:
    _add_heading(doc, title, level=3)
    if counts_df.empty:
        doc.add_paragraph("No data available.")
        return
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = counts_df.columns[0]
    hdr[1].text = "Count"
    for row in counts_df.itertuples(index=False):
        cells = table.add_row().cells
        cells[0].text = str(row[0])
        cells[1].text = str(row[1])
    doc.add_paragraph("")


def _add_picture(doc: Document, chart_path: Path, width: float = 6.0) -> None:
    if chart_path.exists():
        doc.add_picture(str(chart_path), width=Inches(width))
        doc.add_paragraph("")


def _add_markdown_content(doc: Document, markdown_text: str) -> None:
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)


def build_word_report(
    df: pd.DataFrame,
    charts_dir: Path,
    output_path: Path,
    analysis_text: str = "",
) -> None:
    doc = Document()
    _add_title_page(doc)

    total = len(df)
    bad_fit = int((df["agent_classification"] == "Bad Fit").sum())
    moderate_fit = int((df["agent_classification"] == "Moderate Fit").sum())
    good_fit = int((df["agent_classification"] == "Good Fit").sum())
    avg_conf = df["confidence_score"].mean()
    pass_count = int((df["evaluation_result"] == "Pass").sum())
    pass_obs = int((df["evaluation_result"] == "Pass with Observation").sum())
    pass_fail = int((df["evaluation_result"] == "Fail").sum())
    issued_offered_bad = df[
        df["outcome_segment"].isin(["Issued", "Offered"]) & (df["agent_classification"] == "Bad Fit")
    ]

    _add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph(f"Total tests in this report: {total}")
    doc.add_paragraph(
        f"Classifications — Bad Fit: {bad_fit}; Moderate Fit: {moderate_fit}; Good Fit: {good_fit}"
    )
    doc.add_paragraph(
        f"Average confidence score: {avg_conf:.4f}" if pd.notna(avg_conf) else "Average confidence score: N/A"
    )
    doc.add_paragraph(f"Evaluation — Pass: {pass_count}; Pass with Observation: {pass_obs}; Fail: {pass_fail}")
    doc.add_paragraph(
        f"Issued/Offered records classified as Bad Fit: {len(issued_offered_bad)} "
        "(indicates the agent did not treat historical quote status as underwriting truth)."
    )
    doc.add_paragraph(
        "This report documents segmented underwriting regression testing for the AF-UW-RiskAppetite "
        "agent. Detailed per-test prompts and responses are stored in outputs/agent_test_results.json."
    )

    _add_heading(doc, "Methodology", level=1)
    doc.add_paragraph(
        "The testing sample was created from a Texas segmented CSV generated from the underwriting "
        "query, filtered to Texas and records after 2026-01-01."
    )
    coverage_topics = [
        "VES / Non-Admitted",
        "VRE / Admitted",
        "Issued",
        "Offered",
        "In Progress",
        "Declined",
        "Other Closed / Not Bound",
        "Referred",
        "Closed Coastal County",
        "Harris <= 10 Miles From Coast",
        "Harris / Coastal Sensitivity",
        "Interior / Non-Closed County",
        "Below $1M Coverage A",
        "$1M-$2M Coverage A",
        "$2M-$3M Coverage A",
        "$3M-$10M Coverage A",
        "Missing company/program/location",
        "Protected fire protection",
        "Partial protected fire protection",
        "Unprotected fire protection",
        "Alarm and sprinkler variations",
        "Data quality issues",
    ]
    doc.add_paragraph("Prompts were selected to cover:")
    for topic in coverage_topics:
        doc.add_paragraph(topic, style="List Bullet")

    _add_heading(doc, "Data Segmentation Approach", level=1)
    segmentation_fields = [
        "guideline_track",
        "outcome_segment",
        "close_reason_segment",
        "territory_segment",
        "coverage_segment",
        "tiv_segment",
        "loss_segment",
        "alarm_segment",
        "sprinkler_segment",
        "gated_segment",
        "protection_class_segment",
        "fire_protection_segment",
        "expected_agent_focus",
    ]
    doc.add_paragraph("The input CSV included these segmentation fields:")
    for field in segmentation_fields:
        doc.add_paragraph(field, style="List Bullet")
    doc.add_paragraph(
        "The `expected_agent_focus` field was kept only for local evaluation and was not sent to the "
        "agent to avoid biasing the response."
    )

    _add_heading(doc, "Segment Record Counts", level=1)
    for column, title in [
        ("guideline_track", "Guideline Track"),
        ("outcome_segment", "Outcome Segment"),
        ("territory_segment", "Territory Segment"),
        ("coverage_segment", "Coverage Segment"),
        ("tiv_segment", "TIV Segment"),
        ("alarm_segment", "Alarm Segment"),
        ("sprinkler_segment", "Sprinkler Segment"),
        ("fire_protection_segment", "Fire Protection Segment"),
        ("agent_classification", "Agent Classification"),
        ("evaluation_result", "Evaluation Result"),
    ]:
        _add_counts_table(doc, _count_table(df, column), title)

    _add_heading(doc, "Charts", level=1)
    chart_titles = {
        "01_classification_distribution.png": "Classification Distribution",
        "02_evaluation_distribution.png": "Evaluation Distribution",
        "03_confidence_by_test.png": "Confidence Score by Test",
        "04_classification_by_guideline_track.png": "Classification by Guideline Track",
        "05_classification_by_outcome_segment.png": "Classification by Outcome Segment",
        "06_classification_by_territory_segment.png": "Classification by Territory Segment",
        "07_average_confidence_by_segment.png": "Average Confidence by Major Segment",
        "08_missing_information_frequency.png": "Missing Information Frequency",
    }
    for filename, title in chart_titles.items():
        _add_heading(doc, title, level=2)
        _add_picture(doc, charts_dir / filename)

    _add_heading(doc, "Results Analysis & Conclusions", level=1)
    if analysis_text.strip():
        body = analysis_text
        if body.startswith("# Results Analysis"):
            body = "\n".join(body.splitlines()[2:])
        _add_markdown_content(doc, body)
    else:
        doc.add_paragraph(
            "Analysis not available. See outputs/results_analysis.md after report generation completes."
        )

    _add_heading(doc, "Data Reference", level=1)
    doc.add_paragraph(
        "Per-test prompts, submissions, and full agent JSON responses are available in "
        "outputs/agent_test_results.json and outputs/agent_test_results.csv. "
        "This Word report intentionally excludes per-test detail to keep the document concise."
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def generate_reports(
    results_csv_path: str | Path,
    output_dir: str | Path,
    planned_total: int | None = None,
) -> dict[str, Path]:
    """Create Markdown analysis, charts, and Word report from agent test results."""
    from report_analyst import generate_results_analysis

    results_csv_path = Path(results_csv_path)
    output_dir = Path(output_dir)
    charts_dir = output_dir / "charts"

    if not results_csv_path.exists():
        raise FileNotFoundError(f"Results CSV not found: {results_csv_path}")

    df = enrich_results(pd.read_csv(results_csv_path))
    chart_paths = generate_charts(df, charts_dir)

    markdown_path = output_dir / "segment_analysis.md"
    markdown_path.write_text(build_segment_analysis_markdown(df), encoding="utf-8")

    analysis_text = generate_results_analysis(df, output_dir, planned_total=planned_total)
    analysis_path = output_dir / "results_analysis.md"

    docx_path = output_dir / REPORT_DOCX_NAME
    build_word_report(df, charts_dir, docx_path, analysis_text=analysis_text)

    return {
        "segment_analysis_md": markdown_path,
        "results_analysis_md": analysis_path,
        "report_docx": docx_path,
        **chart_paths,
    }
